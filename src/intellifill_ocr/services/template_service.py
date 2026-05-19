from __future__ import annotations

from pathlib import Path

import pandas as pd
import pdfplumber
from docx import Document
from openpyxl import load_workbook

from intellifill_ocr.models.template import TemplateCell, TemplateTable
from intellifill_ocr.ocr.engine import OCREngine
from intellifill_ocr.ocr.pdf import parse_pdf_text
from intellifill_ocr.utils.placeholders import is_placeholder_text


class TemplateService:
    def __init__(self, ocr_engine: OCREngine) -> None:
        self.ocr_engine = ocr_engine

    def load_template(self, path: Path) -> TemplateTable:
        suffix = path.suffix.lower()
        if suffix == ".xlsx":
            return self._from_xlsx(path)
        if suffix == ".xls":
            return self._from_excel_workbook(path)
        if suffix == ".csv":
            return self._from_dataframe(path.stem, pd.read_csv(path, header=None, dtype=str).fillna(""))
        if suffix == ".docx":
            return self._from_docx(path)
        if suffix == ".pdf":
            return self._from_pdf(path)
        if suffix in {".png", ".jpg", ".jpeg"}:
            text, _ = self.ocr_engine.image_to_text(path)
            rows = [
                [TemplateCell(row=i, column=0, value=line, table_index=0)]
                for i, line in enumerate(text.splitlines())
                if line.strip()
            ]
            return self._attach_document_tables(path.stem, [TemplateTable(name=path.stem, cells=rows)])
        raise ValueError(f"Unsupported template type: {path.suffix}")

    def _attach_document_tables(self, name: str, tables: list[TemplateTable]) -> TemplateTable:
        if not tables:
            tables = [TemplateTable(name=name)]
        for table_index, table in enumerate(tables):
            table.name = name
            table.table_index = table_index
            if not table.display_name:
                table.display_name = f"Table {table_index + 1}"
            for row_index, row in enumerate(table.cells):
                for column_index, cell in enumerate(row):
                    cell.row = row_index
                    cell.column = column_index
                    cell.table_index = table_index
            table.document_tables = tables
        return tables[0]

    def _from_dataframe(
        self,
        name: str,
        frame: pd.DataFrame,
        table_index: int = 0,
        display_name: str = "",
    ) -> TemplateTable:
        rows: list[list[TemplateCell]] = []
        for row_index, row in frame.iterrows():
            cells: list[TemplateCell] = []
            for col_index, value in enumerate(row.tolist()):
                text = "" if pd.isna(value) else str(value)
                cells.append(
                    TemplateCell(
                        row=row_index,
                        column=col_index,
                        value=text,
                        is_placeholder=is_placeholder_text(text),
                        table_index=table_index,
                    )
                )
            rows.append(cells)
        return TemplateTable(name=name, cells=rows, table_index=table_index, display_name=display_name)

    def _from_xlsx(self, path: Path) -> TemplateTable:
        workbook = load_workbook(path)
        tables = [
            self._from_worksheet(path.stem, sheet, table_index, sheet.title)
            for table_index, sheet in enumerate(workbook.worksheets)
        ]
        return self._attach_document_tables(path.stem, tables)

    def _from_worksheet(self, name: str, sheet, table_index: int, display_name: str) -> TemplateTable:
        max_row = sheet.max_row or 1
        max_col = sheet.max_column or 1
        span_map: dict[tuple[int, int], tuple[int, int]] = {}
        covered: set[tuple[int, int]] = set()
        for merged in sheet.merged_cells.ranges:
            min_col, min_row, max_col_range, max_row_range = merged.bounds
            span_map[(min_row - 1, min_col - 1)] = (max_row_range - min_row + 1, max_col_range - min_col + 1)
            for row in range(min_row - 1, max_row_range):
                for col in range(min_col - 1, max_col_range):
                    if (row, col) != (min_row - 1, min_col - 1):
                        covered.add((row, col))

        rows: list[list[TemplateCell]] = []
        for row_index in range(max_row):
            row_cells: list[TemplateCell] = []
            for col_index in range(max_col):
                cell_value = sheet.cell(row=row_index + 1, column=col_index + 1).value
                text = "" if cell_value is None else str(cell_value)
                row_span, column_span = span_map.get((row_index, col_index), (1, 1))
                row_cells.append(
                    TemplateCell(
                        row=row_index,
                        column=col_index,
                        value=text,
                        is_placeholder=is_placeholder_text(text) or (row_index, col_index) in covered,
                        row_span=row_span,
                        column_span=column_span,
                        table_index=table_index,
                    )
                )
            rows.append(row_cells)
        return TemplateTable(name=name, cells=rows, table_index=table_index, display_name=display_name)

    def _from_excel_workbook(self, path: Path) -> TemplateTable:
        sheets = pd.read_excel(path, sheet_name=None, header=None, dtype=str)
        tables = [
            self._from_dataframe(path.stem, frame.fillna(""), table_index, str(sheet_name))
            for table_index, (sheet_name, frame) in enumerate(sheets.items())
        ]
        return self._attach_document_tables(path.stem, tables)

    def _from_pdf(self, path: Path) -> TemplateTable:
        tables_out: list[TemplateTable] = []
        with pdfplumber.open(path) as pdf:
            for page_index, page in enumerate(pdf.pages):
                tables = page.find_tables()
                if not tables:
                    continue
                for page_table_index, table in enumerate(tables):
                    table_index = len(tables_out)
                    extracted = table.extract()
                    rows: list[list[TemplateCell]] = []
                    for row_index, row in enumerate(extracted):
                        row_cells: list[TemplateCell] = []
                        table_row = table.rows[row_index] if row_index < len(table.rows) else None
                        cell_boxes = table_row.cells if table_row else []
                        for col_index, value in enumerate(row):
                            text = (value or "").strip()
                            bbox = cell_boxes[col_index] if col_index < len(cell_boxes) else None
                            row_cells.append(
                                TemplateCell(
                                    row=row_index,
                                    column=col_index,
                                    value=text,
                                    is_placeholder=is_placeholder_text(text),
                                    table_index=table_index,
                                    source_page=page_index,
                                    bbox=tuple(float(part) for part in bbox) if bbox else None,
                                )
                            )
                        rows.append(row_cells)
                    if rows:
                        tables_out.append(
                            TemplateTable(
                                name=path.stem,
                                cells=rows,
                                table_index=table_index,
                                display_name=f"Page {page_index + 1} Table {page_table_index + 1}",
                            )
                        )
        if tables_out:
            return self._attach_document_tables(path.stem, tables_out)

        parsed = parse_pdf_text(path, self.ocr_engine)
        if parsed.tables:
            tables = []
            for table_index, table in enumerate(parsed.tables):
                rows: list[list[TemplateCell]] = []
                for row_index, row in enumerate(table):
                    rows.append(
                        [
                            TemplateCell(
                                row=row_index,
                                column=col_index,
                                value=value,
                                is_placeholder=is_placeholder_text(value),
                                table_index=table_index,
                            )
                            for col_index, value in enumerate(row)
                        ]
                    )
                tables.append(
                    TemplateTable(
                        name=path.stem,
                        cells=rows,
                        table_index=table_index,
                        display_name=f"Table {table_index + 1}",
                    )
                )
            return self._attach_document_tables(path.stem, tables)

        rows = [
            [TemplateCell(row=i, column=0, value=line, table_index=0)]
            for i, line in enumerate(parsed.text.splitlines())
            if line.strip()
        ]
        return self._attach_document_tables(path.stem, [TemplateTable(name=path.stem, cells=rows)])

    def _from_docx(self, path: Path) -> TemplateTable:
        doc = Document(path)
        doc_tables = list(self._iter_docx_tables(doc))
        if not doc_tables:
            rows = [
                [TemplateCell(row=i, column=0, value=p.text, table_index=0)]
                for i, p in enumerate(doc.paragraphs)
                if p.text.strip()
            ]
            return self._attach_document_tables(path.stem, [TemplateTable(name=path.stem, cells=rows)])

        tables = [
            self._from_docx_table(path.stem, table, table_index)
            for table_index, table in enumerate(doc_tables)
        ]
        return self._attach_document_tables(path.stem, tables)

    def _iter_docx_tables(self, container):
        for table in container.tables:
            yield table
            for row in table.rows:
                for cell in row.cells:
                    yield from self._iter_docx_tables(cell)

    def _from_docx_table(self, name: str, table, table_index: int) -> TemplateTable:
        rows: list[list[TemplateCell]] = []
        for row_index, row in enumerate(table.rows):
            cells: list[TemplateCell] = []
            for col_index, cell in enumerate(row.cells):
                text = cell.text.strip()
                cells.append(
                    TemplateCell(
                        row=row_index,
                        column=col_index,
                        value=text,
                        is_placeholder=is_placeholder_text(text),
                        table_index=table_index,
                    )
                )
            rows.append(cells)
        return TemplateTable(name=name, cells=rows, table_index=table_index, display_name=f"Table {table_index + 1}")
