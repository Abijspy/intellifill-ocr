from __future__ import annotations

from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches
from openpyxl import load_workbook
from openpyxl.cell.cell import MergedCell
from openpyxl.drawing.image import Image as ExcelImage

from intellifill_ocr.models.template import TemplateCell
from intellifill_ocr.models.template import TemplateTable
from intellifill_ocr.utils.barcode import barcode_png_bytes
from intellifill_ocr.utils.placeholders import is_placeholder_text


class PreservedTemplateExporter:
    def export(self, template_path: Path, table: TemplateTable, output_path: Path, traceability_code: str = "") -> None:
        suffix = template_path.suffix.lower()
        if suffix == ".xlsx":
            self._export_xlsx(template_path, table, output_path, traceability_code)
            return
        if suffix == ".docx":
            self._export_docx(template_path, table, output_path, traceability_code)
            return
        if suffix == ".csv":
            rows: list[list[str]] = []
            for template_table in table.all_tables():
                if rows:
                    rows.append([])
                if table.table_count > 1:
                    rows.append([template_table.label])
                rows.extend([[cell.value for cell in row] for row in template_table.cells])
            if traceability_code:
                rows.extend([[], ["Traceability ID", traceability_code]])
            pd.DataFrame(rows).to_csv(output_path, index=False, header=False)
            return
        raise ValueError("Original-format export is currently available for DOCX, XLSX, and CSV templates.")

    def _export_xlsx(self, template_path: Path, table: TemplateTable, output_path: Path, traceability_code: str) -> None:
        workbook = load_workbook(template_path)
        for template_table in table.all_tables():
            if template_table.table_index >= len(workbook.worksheets):
                continue
            sheet = workbook.worksheets[template_table.table_index]
            for row_index, row in enumerate(template_table.cells, start=1):
                for col_index, cell in enumerate(row, start=1):
                    target_cell = sheet.cell(row=row_index, column=col_index)
                    if isinstance(target_cell, MergedCell):
                        continue
                    if self._should_fill_preserved_cell(cell, target_cell.value):
                        target_cell.value = cell.value
        if traceability_code:
            sheet = workbook.active
            traceability_row = sheet.max_row + 2
            sheet.cell(row=traceability_row, column=1).value = "Traceability ID"
            sheet.cell(row=traceability_row, column=2).value = traceability_code
            image = ExcelImage(barcode_png_bytes(traceability_code))
            image.width = 255
            image.height = 58
            sheet.add_image(image, f"A{traceability_row + 1}")
        workbook.save(output_path)

    def _export_docx(self, template_path: Path, table: TemplateTable, output_path: Path, traceability_code: str) -> None:
        document = Document(template_path)
        document_tables = list(self._iter_docx_tables(document))
        if not document_tables:
            self._append_docx_traceability(document, traceability_code)
            document.save(output_path)
            return

        for template_table in table.all_tables():
            if template_table.table_index >= len(document_tables):
                continue
            target_table = document_tables[template_table.table_index]
            for row_index, row in enumerate(template_table.cells):
                if row_index >= len(target_table.rows):
                    break
                for col_index, source_cell in enumerate(row):
                    if col_index >= len(target_table.rows[row_index].cells):
                        break
                    target_cell = target_table.rows[row_index].cells[col_index]
                    if self._should_fill_preserved_cell(source_cell, self._cell_text(target_cell)):
                        self._set_cell_text(target_cell, source_cell.value)
        self._append_docx_traceability(document, traceability_code)
        document.save(output_path)

    def _iter_docx_tables(self, container):
        for table in container.tables:
            yield table
            for row in table.rows:
                for cell in row.cells:
                    yield from self._iter_docx_tables(cell)

    def _should_fill_preserved_cell(self, cell: TemplateCell, original_value: object | None) -> bool:
        return cell.is_placeholder and bool(cell.value.strip()) and is_placeholder_text(original_value)

    def _cell_text(self, cell) -> str:
        return "\n".join(paragraph.text for paragraph in cell.paragraphs).strip()

    def _set_cell_text(self, cell, value: str) -> None:
        paragraph = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
        if paragraph.runs:
            paragraph.runs[0].text = value
            for run in paragraph.runs[1:]:
                run.text = ""
        else:
            paragraph.add_run(value)

    def _append_docx_traceability(self, document: Document, traceability_code: str) -> None:
        if not traceability_code:
            return
        footer = document.sections[0].footer
        paragraph = footer.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.add_run(f"Traceability ID: {traceability_code}")
        paragraph.add_run().add_break()
        paragraph.add_run().add_picture(
            barcode_png_bytes(traceability_code, narrow=1, bar_height=30),
            width=Inches(2.55),
        )
