from __future__ import annotations

import tempfile
from pathlib import Path

import pandas as pd
import pdfplumber
import pypdfium2 as pdfium
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches
from PIL import Image, ImageDraw, ImageFont

from intellifill_ocr.models.template import TemplateTable
from intellifill_ocr.utils.barcode import barcode_png_bytes


class ExportService:
    def to_dataframe(self, template: TemplateTable) -> pd.DataFrame:
        rows = [[cell.value for cell in row] for row in template.cells]
        return pd.DataFrame(rows)

    def export_csv(self, template: TemplateTable, path: Path) -> None:
        tables = template.all_tables()
        if len(tables) == 1:
            self.to_dataframe(tables[0]).to_csv(path, header=False, index=False)
            return

        rows: list[list[str]] = []
        for table in tables:
            if rows:
                rows.append([])
            rows.append([table.label])
            rows.extend([[cell.value for cell in row] for row in table.cells])
        pd.DataFrame(rows).to_csv(path, header=False, index=False)

    def export_excel(self, template: TemplateTable, path: Path) -> None:
        with pd.ExcelWriter(path, engine="openpyxl") as writer:
            tables = template.all_tables()
            for table in tables:
                sheet_name = self._excel_sheet_name(table.label, table.table_index)
                self.to_dataframe(table).to_excel(writer, sheet_name=sheet_name, header=False, index=False)

    def export_word(self, template: TemplateTable, path: Path, traceability_code: str = "") -> None:
        document = Document()
        document.add_heading(template.name or "Filled Template", level=1)
        tables = template.all_tables()
        for table_index, template_table in enumerate(tables):
            if len(tables) > 1:
                if table_index:
                    document.add_paragraph()
                document.add_heading(template_table.label, level=2)
            table = document.add_table(rows=max(template_table.row_count, 1), cols=max(template_table.column_count, 1))
            table.style = "Table Grid"
            for row_index, row in enumerate(template_table.cells):
                for column_index in range(template_table.column_count):
                    value = row[column_index].value if column_index < len(row) else ""
                    table.cell(row_index, column_index).text = value
        self._append_word_traceability(document, traceability_code)
        document.save(path)

    def export_pdf(self, template: TemplateTable, path: Path, traceability_code: str = "") -> None:
        pages: list[Image.Image] = []
        page = self._new_pdf_page()
        draw = ImageDraw.Draw(page)
        body_font = self._font(11)
        header_font = self._font(13, bold=True)
        y = 45
        row_height = 28
        reserved_footer = 104 if traceability_code else 45

        for table_index, template_table in enumerate(template.all_tables()):
            if table_index:
                y += 28
            if y > page.height - reserved_footer - 60:
                self._draw_traceability_footer(page, traceability_code)
                pages.append(page)
                page = self._new_pdf_page()
                draw = ImageDraw.Draw(page)
                y = 45

            if template.table_count > 1:
                draw.text((40, y), template_table.label, fill="#111827", font=header_font)
                y += 28

            col_width = max(90, int((page.width - 80) / max(template_table.column_count, 1)))
            for row in template_table.cells:
                x = 40
                for cell in row:
                    rect = (x, y, x + col_width, y + row_height)
                    draw.rectangle(rect, outline="#111827", width=1)
                    self._draw_text_in_rect(draw, cell.value, rect, body_font)
                    x += col_width
                y += row_height
                if y > page.height - reserved_footer:
                    self._draw_traceability_footer(page, traceability_code)
                    pages.append(page)
                    page = self._new_pdf_page()
                    draw = ImageDraw.Draw(page)
                    y = 45

        self._draw_traceability_footer(page, traceability_code)
        pages.append(page)
        self._save_pdf_pages(pages, path)

    def export_preserved_pdf(
        self,
        template_path: Path,
        template: TemplateTable,
        path: Path,
        traceability_code: str = "",
    ) -> None:
        if template_path.suffix.lower() != ".pdf" or not self._has_pdf_coordinates(template):
            self.export_pdf(template, path, traceability_code)
            return

        pages: list[Image.Image] = []
        overlays = 0
        rendered_pdf = pdfium.PdfDocument(str(template_path))
        with pdfplumber.open(template_path) as source_pdf, tempfile.TemporaryDirectory(prefix="intellifill_pdf_export_") as temp_dir:
            for page_index, source_page in enumerate(source_pdf.pages):
                bitmap = rendered_pdf[page_index].render(scale=2.0).to_pil().convert("RGB")
                image_path = Path(temp_dir) / f"page_{page_index + 1}.png"
                bitmap.save(image_path)
                image = Image.open(image_path).convert("RGB")
                draw = ImageDraw.Draw(image)
                font = self._font(18)

                if page_index == 0:
                    self._draw_traceability_footer(image, traceability_code)

                scale_x = image.width / max(float(source_page.width), 1.0)
                scale_y = image.height / max(float(source_page.height), 1.0)
                for table in template.all_tables():
                    for row in table.cells:
                        for cell in row:
                            if not cell.is_placeholder or not cell.value.strip() or cell.source_page != page_index or not cell.bbox:
                                continue
                            x0, top, x1, bottom = cell.bbox
                            rect = (
                                int(x0 * scale_x) + 6,
                                int(top * scale_y) + 4,
                                int(x1 * scale_x) - 6,
                                int(bottom * scale_y) - 4,
                            )
                            self._draw_text_in_rect(draw, cell.value, rect, font)
                            overlays += 1
                pages.append(image)

        if overlays == 0:
            self.export_pdf(template, path, traceability_code)
            return

        self._save_pdf_pages(pages, path)

    def _draw_traceability_footer(self, page: Image.Image, traceability_code: str) -> None:
        if not traceability_code:
            return
        draw = ImageDraw.Draw(page)
        barcode = Image.open(barcode_png_bytes(traceability_code, narrow=2, bar_height=34, show_text=False, quiet_zone=8)).convert("RGB")
        max_block_width = max(220, page.width - 80)
        barcode_width = min(barcode.width, max_block_width - 32)
        barcode_height = max(28, int(barcode.height * (barcode_width / max(barcode.width, 1))))
        barcode = barcode.resize((barcode_width, barcode_height), Image.Resampling.NEAREST)
        block_width = barcode_width + 32
        block_height = 22 + barcode_height + 16
        x = int((page.width - block_width) / 2)
        y = page.height - block_height - 12
        draw.rectangle((x - 6, y - 6, x + block_width + 6, y + block_height + 6), fill="white")
        label = f"Traceability ID: {traceability_code}"
        label_font = self._font(10)
        label_width = int(draw.textlength(label, font=label_font))
        draw.text((x + int((block_width - label_width) / 2), y), label, fill="#111827", font=label_font)
        page.paste(barcode, (x + 16, y + 22))

    def _append_word_traceability(self, document: Document, traceability_code: str) -> None:
        if not traceability_code:
            return
        footer = document.sections[0].footer
        paragraph = footer.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.add_run(f"Traceability ID: {traceability_code}")
        paragraph.add_run().add_break()
        paragraph.add_run().add_picture(
            barcode_png_bytes(traceability_code, narrow=1, bar_height=30),
            width=Inches(2.55),
        )

    def _has_pdf_coordinates(self, template: TemplateTable) -> bool:
        return any(cell.bbox for table in template.all_tables() for row in table.cells for cell in row)

    def _excel_sheet_name(self, label: str, table_index: int) -> str:
        cleaned = "".join("_" if char in "[]:*?/\\" else char for char in label).strip()
        cleaned = cleaned or f"Table {table_index + 1}"
        return cleaned[:31]

    @staticmethod
    def _new_pdf_page() -> Image.Image:
        return Image.new("RGB", (794, 1123), "white")

    @staticmethod
    def _font(size: int, bold: bool = False) -> ImageFont.ImageFont:
        font_names = ["arialbd.ttf" if bold else "arial.ttf", "seguisb.ttf" if bold else "segoeui.ttf"]
        for font_name in font_names:
            try:
                return ImageFont.truetype(font_name, size)
            except OSError:
                continue
        return ImageFont.load_default()

    @staticmethod
    def _draw_text_in_rect(
        draw: ImageDraw.ImageDraw,
        text: str,
        rect: tuple[int, int, int, int],
        font: ImageFont.ImageFont,
    ) -> None:
        x0, y0, x1, y1 = rect
        words = str(text or "").split()
        if not words:
            return
        lines: list[str] = []
        current = ""
        max_width = max(8, x1 - x0 - 8)
        for word in words:
            candidate = f"{current} {word}".strip()
            if draw.textlength(candidate, font=font) <= max_width or not current:
                current = candidate
            else:
                lines.append(current)
                current = word
        if current:
            lines.append(current)
        line_height = max(12, int(font.size * 1.25) if hasattr(font, "size") else 14)
        y = y0 + 4
        for line in lines:
            if y + line_height > y1:
                break
            draw.text((x0 + 4, y), line, fill="#111827", font=font)
            y += line_height

    @staticmethod
    def _save_pdf_pages(pages: list[Image.Image], path: Path) -> None:
        if not pages:
            pages = [ExportService._new_pdf_page()]
        output_pages = [page.convert("RGB") for page in pages]
        first, rest = output_pages[0], output_pages[1:]
        first.save(path, "PDF", resolution=96.0, save_all=bool(rest), append_images=rest)
