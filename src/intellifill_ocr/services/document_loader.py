from __future__ import annotations

import csv
from pathlib import Path

import pandas as pd
from docx import Document

from intellifill_ocr.models.document import ParsedDocument
from intellifill_ocr.ocr.engine import OCREngine
from intellifill_ocr.ocr.pdf import parse_pdf_text
from intellifill_ocr.utils.exceptions import UnsupportedDocumentError


class DocumentLoader:
    SUPPORTED = {".docx", ".xlsx", ".xls", ".csv", ".png", ".jpg", ".jpeg", ".pdf"}

    def __init__(self, ocr_engine: OCREngine) -> None:
        self.ocr_engine = ocr_engine

    def parse(self, path: Path) -> ParsedDocument:
        suffix = path.suffix.lower()
        if suffix not in self.SUPPORTED:
            raise UnsupportedDocumentError(f"Unsupported file type: {path.suffix}")
        if suffix == ".docx":
            return self._parse_docx(path)
        if suffix in {".xlsx", ".xls"}:
            return self._parse_excel(path)
        if suffix == ".csv":
            return self._parse_csv(path)
        if suffix in {".png", ".jpg", ".jpeg"}:
            text, boxes = self.ocr_engine.image_to_text(path)
            return ParsedDocument(path=path, text=text, ocr_boxes=boxes)
        if suffix == ".pdf":
            return parse_pdf_text(path, self.ocr_engine)
        raise UnsupportedDocumentError(f"Unsupported file type: {path.suffix}")

    def _parse_docx(self, path: Path) -> ParsedDocument:
        doc = Document(path)
        text_parts = [p.text for p in doc.paragraphs if p.text.strip()]
        tables: list[list[list[str]]] = []
        for table in self._iter_docx_tables(doc):
            parsed_table: list[list[str]] = []
            for row in table.rows:
                parsed_table.append([cell.text.strip() for cell in row.cells])
            tables.append(parsed_table)
            text_parts.extend(" | ".join(row) for row in parsed_table)
        return ParsedDocument(path=path, text="\n".join(text_parts), tables=tables)

    def _iter_docx_tables(self, container):
        for table in container.tables:
            yield table
            for row in table.rows:
                for cell in row.cells:
                    yield from self._iter_docx_tables(cell)

    def _parse_excel(self, path: Path) -> ParsedDocument:
        sheets = pd.read_excel(path, sheet_name=None, header=None, dtype=str)
        text_parts: list[str] = []
        tables: list[list[list[str]]] = []
        for sheet_name, frame in sheets.items():
            clean = frame.fillna("")
            table = clean.astype(str).values.tolist()
            tables.append(table)
            text_parts.append(f"[{sheet_name}]")
            text_parts.extend(" | ".join(row) for row in table)
        return ParsedDocument(path=path, text="\n".join(text_parts), tables=tables)

    def _parse_csv(self, path: Path) -> ParsedDocument:
        rows: list[list[str]] = []
        with path.open("r", encoding="utf-8-sig", newline="") as handle:
            for row in csv.reader(handle):
                rows.append([cell.strip() for cell in row])
        expected_columns = len(next((row for row in rows if row), []))
        if expected_columns == 2:
            rows = [
                [row[0], ", ".join(cell.strip() for cell in row[1:])] if len(row) > 2 else row
                for row in rows
            ]
        max_columns = max((len(row) for row in rows), default=0)
        table = [row + [""] * (max_columns - len(row)) for row in rows]
        text = "\n".join(" | ".join(row) for row in table)
        return ParsedDocument(path=path, text=text, tables=[table])
