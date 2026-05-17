from pathlib import Path

from docx import Document
from openpyxl import Workbook

ROOT = Path(__file__).resolve().parent


def create_docx() -> None:
    doc = Document()
    doc.add_heading("Invoice Demo Source", level=1)
    table = doc.add_table(rows=7, cols=2)
    data = [
        ("Invoice No", "INV-2026-1042"),
        ("Date", "2026-05-16"),
        ("Cust Name", "Aster Manufacturing Pvt Ltd"),
        ("Address", "42 Market Street, Pune"),
        ("Subtotal", "12500.00"),
        ("GST", "2250.00"),
        ("Grand Total", "14750.00"),
    ]
    for row, (key, value) in zip(table.rows, data):
        row.cells[0].text = key
        row.cells[1].text = value
    doc.save(ROOT / "source_invoice.docx")


def create_xlsx() -> None:
    wb = Workbook()
    ws = wb.active
    ws.title = "Invoice"
    rows = [
        ("Field", "Value", "Notes"),
        ("Invoice Number", "", ""),
        ("Invoice Date", "", ""),
        ("Customer Name", "", ""),
        ("Customer Address", "", ""),
        ("Subtotal", "", ""),
        ("Tax", "", ""),
        ("Total Amount", "", ""),
    ]
    for row in rows:
        ws.append(row)
    wb.save(ROOT / "template_invoice.xlsx")


if __name__ == "__main__":
    create_docx()
    create_xlsx()
    print("Created demo/source_invoice.docx and demo/template_invoice.xlsx")
